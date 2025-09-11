import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
import io

def pdf_to_word(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    st.write("📄 Number of pages in PDF:", len(pdf_reader.pages))

    doc = Document()
    for i, page in enumerate(pdf_reader.pages, start=1):
        text = page.extract_text()
        if text:  # Only add if text exists
            doc.add_paragraph(text)
        else:
            doc.add_paragraph(f"[Page {i} contains no extractable text]")
    return doc

def main():
    st.title("📑 PDF to Word Converter")
    uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")

    if uploaded_file is not None:
        doc = pdf_to_word(uploaded_file)
        st.success("✅ PDF converted to Word document successfully!")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 Download Word Document",
            data=buffer,
            file_name="converted.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
